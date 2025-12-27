import os
import torch
import re
from flask import Flask, request, jsonify, render_template, send_from_directory, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename
from PIL import Image
import io
import base64
from config import *
from transformers import AutoModel, AutoTokenizer

def format_ocr_result(text, output_format='markdown'):
    """Format OCR result - Giữ nguyên output từ DeepSeek-OCR"""
    if not text:
        return text
    
    # Nếu output đã là markdown sạch (không có tags), giữ nguyên 100%
    if output_format == 'markdown' and '<|ref|>' not in text and '<|det|>' not in text:
        # Đã là markdown sạch từ model, giữ nguyên hoàn toàn
        return text
    
    # Nếu có <|ref|> và <|det|> tags (grounding format)
    if '<|ref|>' in text and '<|det|>' in text:
        if output_format == 'markdown':
            # Convert sang markdown đẹp
            return format_to_markdown(text)
        elif output_format == 'full':
            # Giữ nguyên format với bounding boxes
            return format_with_boxes(text)
        else:
            # Chỉ lấy text, bỏ tags
            return extract_text_only(text)
    else:
        # Không có tags, trả về nguyên bản (có thể đã là markdown)
        return text

def format_to_markdown(text):
    """Convert OCR result với <|ref|> tags sang markdown - Giữ layout đẹp"""
    lines = text.split('\n')
    markdown_lines = []
    last_tag = None
    last_was_title = False
    
    for i, line in enumerate(lines):
        original_line = line
        line_stripped = line.strip()
        
        # Giữ empty lines để giữ layout
        if not line_stripped:
            if last_was_title:
                continue  # Bỏ empty line sau title
            markdown_lines.append('')
            continue
            
        # Parse <|ref|>tag<|/ref|><|det|>[[x1,y1,x2,y2]]<|/det|>
        ref_match = re.search(r'<\|ref\|>(.*?)<\|/ref\|>', line)
        det_match = re.search(r'<\|det\|>\[\[(.*?)\]\]<\|/det\|>', line)
        
        if ref_match:
            tag = ref_match.group(1)
            # Lấy text sau tags, giữ nguyên spacing
            text_part = re.sub(r'<\|ref\|>.*?<\|/ref\|>', '', line)
            text_part = re.sub(r'<\|det\|>.*?<\|/det\|>', '', text_part)
            text_part = text_part.strip()
            
            if not text_part:
                continue
            
            # Format theo tag type với layout đẹp
            if tag in ['sub_title', 'title', 'heading']:
                # Thêm spacing trước title
                if not last_was_title and markdown_lines and markdown_lines[-1]:
                    markdown_lines.append('')
                markdown_lines.append(f'## {text_part}')
                markdown_lines.append('')  # Empty line sau title
                last_was_title = True
                last_tag = tag
            elif tag == 'text':
                # Text bình thường, giữ nguyên
                markdown_lines.append(text_part)
                last_was_title = False
                last_tag = tag
            elif tag == 'image':
                markdown_lines.append('')
                markdown_lines.append(f'![Image]({text_part})')
                markdown_lines.append('')
                last_was_title = False
            elif tag in ['list_item', 'bullet']:
                markdown_lines.append(f'- {text_part}')
                last_was_title = False
            else:
                # Tag khác, format với bold
                markdown_lines.append(f'**{text_part}**')
                last_was_title = False
                last_tag = tag
        else:
            # Không có tags, giữ nguyên text (có thể là markdown đã được format)
            clean_line = re.sub(r'<\|.*?\|>', '', line).strip()
            if clean_line:
                # Kiểm tra xem có phải markdown không
                if clean_line.startswith('#') or clean_line.startswith('-') or clean_line.startswith('*'):
                    markdown_lines.append(clean_line)
                else:
                    markdown_lines.append(clean_line)
                last_was_title = False
    
    # Clean up: loại bỏ nhiều empty lines liên tiếp
    result = []
    prev_empty = False
    for line in markdown_lines:
        if not line.strip():
            if not prev_empty:
                result.append('')
            prev_empty = True
        else:
            result.append(line)
            prev_empty = False
    
    return '\n'.join(result)

def format_with_boxes(text):
    """Format với bounding boxes info"""
    lines = text.split('\n')
    formatted_lines = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Parse tags
        ref_match = re.search(r'<\|ref\|>(.*?)<\|/ref\|>', line)
        det_match = re.search(r'<\|det\|>\[\[(.*?)\]\]<\|/det\|>', line)
        
        if ref_match and det_match:
            tag = ref_match.group(1)
            bbox = det_match.group(1)
            text_part = re.sub(r'<\|ref\|>.*?<\|/ref\|>', '', line)
            text_part = re.sub(r'<\|det\|>.*?<\|/det\|>', '', text_part).strip()
            
            formatted_lines.append(f'[{tag}] {text_part} | BBox: {bbox}')
        else:
            clean_line = re.sub(r'<\|.*?\|>', '', line).strip()
            if clean_line:
                formatted_lines.append(clean_line)
    
    return '\n'.join(formatted_lines)

def extract_text_only(text):
    """Chỉ lấy text, bỏ tất cả tags"""
    # Remove all tags
    clean_text = re.sub(r'<\|.*?\|>', '', text)
    # Remove bounding boxes format
    clean_text = re.sub(r'\[\[.*?\]\]', '', clean_text)
    # Clean up multiple spaces
    clean_text = re.sub(r'\s+', ' ', clean_text)
    return clean_text.strip()

# Patch để fix lỗi DynamicCache.seen_tokens (transformers >= 4.41)
def patch_dynamic_cache():
    """Patch DynamicCache để tương thích với transformers mới"""
    try:
        from transformers.cache_utils import DynamicCache
        import transformers
        
        # Kiểm tra version transformers
        version_parts = transformers.__version__.split('.')
        major, minor = int(version_parts[0]), int(version_parts[1])
        
        # Nếu transformers >= 4.41, cần patch
        if major > 4 or (major == 4 and minor >= 41):
            # Thêm thuộc tính seen_tokens nếu chưa có
            if not hasattr(DynamicCache, 'seen_tokens'):
                def _get_seen_tokens(self):
                    """Get seen_tokens từ cache_position"""
                    try:
                        if hasattr(self, 'cache_position') and len(self.cache_position) > 0:
                            return len(self.cache_position)
                        elif hasattr(self, 'key_cache') and len(self.key_cache) > 0:
                            # Fallback: tính từ key_cache shape
                            return self.key_cache[0].shape[2] if len(self.key_cache) > 0 else 0
                    except:
                        pass
                    return 0
                
                DynamicCache.seen_tokens = property(_get_seen_tokens)
                
                # Thêm method get_max_length nếu chưa có
                if not hasattr(DynamicCache, 'get_max_length'):
                    def _get_max_length(self):
                        """Get max length từ cache"""
                        try:
                            if hasattr(self, 'get_max_cache_shape'):
                                shape = self.get_max_cache_shape()
                                if shape and len(shape) > 1:
                                    return shape[1]
                        except:
                            pass
                        return None
                    DynamicCache.get_max_length = _get_max_length
                
                print("✅ Đã patch DynamicCache để tương thích với transformers mới")
    except Exception as e:
        print(f"⚠️  Không thể patch DynamicCache: {e}")

# Chạy patch ngay khi import
patch_dynamic_cache()

app = Flask(__name__)
CORS(app)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_FILE_SIZE

# Global model and tokenizer
model = None
tokenizer = None

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def init_model():
    """Initialize the DeepSeek-OCR model"""
    global model, tokenizer
    try:
        # Hiển thị thông tin cấu hình
        print("=" * 60)
        print("Cấu hình hệ thống:")
        print(f"  - Device: {DEVICE}")
        print(f"  - Dtype: {DTYPE}")
        print(f"  - Image size: {IMAGE_SIZE}x{IMAGE_SIZE}")
        print(f"  - Base size: {BASE_SIZE}")
        if torch.cuda.is_available():
            gpu_name = torch.cuda.get_device_name(0)
            gpu_memory = torch.cuda.get_device_properties(0).total_memory / (1024**3)
            print(f"  - GPU: {gpu_name} ({gpu_memory:.1f}GB VRAM)")
        else:
            print("  - GPU: Không có (sử dụng CPU)")
        print("=" * 60)
        
        print("\nĐang tải tokenizer...")
        tokenizer = AutoTokenizer.from_pretrained(
            MODEL_NAME, 
            trust_remote_code=True
        )
        
        print("Đang tải model (có thể mất vài phút lần đầu)...")
        print("Lưu ý: Model sẽ được tải từ Hugging Face (~20-30GB)")
        
        # Kiểm tra và cài flash-attn nếu cần (để có LlamaFlashAttention2)
        has_flash_attn = False
        try:
            import flash_attn
            has_flash_attn = True
            print("✅ flash-attn đã được cài đặt")
        except ImportError:
            print("⚠️  flash-attn chưa được cài.")
            print("   Đang thử tải model với transformers mặc định...")
        
        # Patch để bypass flash attention nếu cần
        if not has_flash_attn:
            try:
                # Thử import từ transformers trước
                from transformers.models.llama import modeling_llama
                if not hasattr(modeling_llama, 'LlamaFlashAttention2'):
                    print("⚠️  LlamaFlashAttention2 không có trong transformers.")
                    print("   Đang tạo workaround...")
                    # Tạo class giả để model code không bị lỗi import
                    class FakeLlamaFlashAttention2:
                        def __init__(self, *args, **kwargs):
                            pass
                    modeling_llama.LlamaFlashAttention2 = FakeLlamaFlashAttention2
                    print("✅ Đã tạo workaround cho flash attention")
            except Exception as e:
                print(f"⚠️  Không thể patch: {e}")
        
        # Thử load model
        try:
            # Không chỉ định _attn_implementation để model tự quyết định
            model = AutoModel.from_pretrained(
                MODEL_NAME,
                trust_remote_code=True,
                use_safetensors=True
            )
        except Exception as e:
            error_msg = str(e)
            if "LlamaFlashAttention2" in error_msg or "flash" in error_msg.lower():
                print("\n" + "="*60)
                print("⚠️  Lỗi liên quan đến flash attention.")
                print("="*60)
                print("\n💡 Giải pháp:")
                print("\n1. Cài wheel và flash-attn:")
                print("   pip install wheel")
                print("   pip install flash-attn==2.7.3 --no-build-isolation")
                print("\n2. Hoặc cập nhật transformers:")
                print("   pip install --upgrade transformers>=4.51.0 accelerate")
                print("\n3. Hoặc cài từ pre-built wheel:")
                print("   pip install flash-attn --no-build-isolation")
                print("="*60)
                raise Exception(f"Model yêu cầu flash-attn hoặc transformers mới hơn. Lỗi: {error_msg}")
            else:
                raise Exception(f"Không thể tải model: {error_msg}")
        
        # Move to device and set dtype
        dtype_map = {
            'bfloat16': torch.bfloat16,
            'float16': torch.float16,
            'float32': torch.float32
        }
        dtype = dtype_map.get(DTYPE, torch.bfloat16)
        
        model = model.eval()
        if torch.cuda.is_available() and DEVICE == 'cuda':
            print(f"Đang chuyển model lên GPU với dtype={DTYPE}...")
            print("⚠️  Quá trình này có thể mất 5-10 phút, vui lòng đợi...")
            print("💡 Đang tải ~6.7GB weights lên GPU...")
            
            # Kiểm tra VRAM trước khi tải
            torch.cuda.empty_cache()
            free_memory = torch.cuda.get_device_properties(0).total_memory - torch.cuda.memory_allocated(0)
            free_memory_gb = free_memory / (1024**3)
            print(f"📊 VRAM còn trống: {free_memory_gb:.1f}GB")
            
            if free_memory_gb < 8:
                print("⚠️  VRAM hơi ít, có thể mất nhiều thời gian hơn...")
            
            # Tải model lên GPU
            try:
                model = model.cuda()
                print("✅ Model đã được chuyển lên GPU")
                print("🔄 Đang chuyển đổi dtype...")
                model = model.to(dtype)
                print("✅ Dtype đã được chuyển đổi")
            except RuntimeError as e:
                if "out of memory" in str(e).lower():
                    print("❌ Lỗi: Hết VRAM!")
                    print("💡 Giải pháp: Giảm IMAGE_SIZE trong config.py hoặc dùng CPU")
                    raise
                else:
                    raise
        else:
            print(f"Đang chuyển model lên CPU với dtype={DTYPE}...")
            print("⚠️  Chạy trên CPU sẽ rất chậm (30-60s/ảnh)...")
            model = model.to(dtype)
        
        # Kiểm tra model đã sẵn sàng
        torch.cuda.empty_cache()
        if torch.cuda.is_available() and DEVICE == 'cuda':
            allocated = torch.cuda.memory_allocated(0) / (1024**3)
            print(f"📊 VRAM đã sử dụng: {allocated:.1f}GB")
        
        print("\n✅ Model đã được tải thành công!")
        print("=" * 60)
        return True
    except Exception as e:
        print(f"\n❌ Lỗi khi tải model: {str(e)}")
        print("\nGợi ý khắc phục:")
        print("  1. Kiểm tra kết nối internet")
        print("  2. Đảm bảo có đủ dung lượng ổ cứng (50GB+)")
        print("  3. Thử đổi DEVICE='cpu' trong config.py nếu GPU có vấn đề")
        print("  4. Giảm IMAGE_SIZE trong config.py nếu thiếu RAM/VRAM")
        return False

@app.route('/')
def index():
    """Render the main page"""
    return render_template('index.html')

@app.route('/health')
def health():
    """Health check endpoint"""
    return jsonify({
        'status': 'ok',
        'model_loaded': model is not None,
        'device': DEVICE,
        'cuda_available': torch.cuda.is_available()
    })

@app.route('/outputs/<path:filename>')
def serve_output(filename):
    """Serve files from output directory (for images in markdown)"""
    try:
        return send_from_directory(OUTPUT_FOLDER, filename)
    except Exception as e:
        return jsonify({'error': str(e)}), 404

@app.route('/api/export-docx', methods=['POST'])
def export_docx():
    """Export markdown to DOCX with images"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import markdown
        from markdown.extensions import fenced_code, tables
        import requests
        from io import BytesIO
        
        data = request.get_json()
        markdown_text = data.get('markdown', '')
        filename = data.get('filename', 'ocr_result.docx')
        
        if not markdown_text:
            return jsonify({'error': 'Không có nội dung markdown'}), 400
        
        # Tạo document
        doc = Document()
        
        # Parse markdown line by line (đơn giản hơn)
        lines = markdown_text.split('\n')
        current_list = None
        
        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                current_list = None
                continue
            
            # Headers
            if line_stripped.startswith('# '):
                doc.add_heading(line_stripped[2:], level=1)
                current_list = None
            elif line_stripped.startswith('## '):
                doc.add_heading(line_stripped[3:], level=2)
                current_list = None
            elif line_stripped.startswith('### '):
                doc.add_heading(line_stripped[4:], level=3)
                current_list = None
            elif line_stripped.startswith('#### '):
                doc.add_heading(line_stripped[5:], level=4)
                current_list = None
            # Images
            elif '![' in line_stripped:
                img_match = re.search(r'!\[([^\]]*)\]\(([^\)]+)\)', line_stripped)
                if img_match:
                    img_path = img_match.group(2)
                    if img_path.startswith('/outputs/'):
                        img_path = os.path.join(OUTPUT_FOLDER, img_path.replace('/outputs/', ''))
                    elif not os.path.isabs(img_path):
                        img_path = os.path.join(OUTPUT_FOLDER, img_path)
                    
                    if os.path.exists(img_path):
                        para = doc.add_paragraph()
                        run = para.add_run()
                        try:
                            run.add_picture(img_path, width=Inches(5))
                        except Exception as e:
                            print(f"Error adding image {img_path}: {e}")
                    else:
                        doc.add_paragraph(f"[Image: {img_match.group(1)}]")
                current_list = None
            # Lists
            elif line_stripped.startswith('- ') or line_stripped.startswith('* '):
                if current_list != 'ul':
                    current_list = 'ul'
                para = doc.add_paragraph(line_stripped[2:], style='List Bullet')
            elif re.match(r'^\d+\.\s', line_stripped):
                if current_list != 'ol':
                    current_list = 'ol'
                para = doc.add_paragraph(re.sub(r'^\d+\.\s', '', line_stripped), style='List Number')
            # Code blocks
            elif line_stripped.startswith('```'):
                continue
            # Regular text
            else:
                # Remove markdown formatting
                text = re.sub(r'\*\*(.*?)\*\*', r'\1', line_stripped)  # Bold
                text = re.sub(r'\*(.*?)\*', r'\1', text)  # Italic
                text = re.sub(r'`([^`]+)`', r'\1', text)  # Code
                para = doc.add_paragraph(text)
                current_list = None
        
        # Skip HTML parser, use simple markdown parsing instead
        from html.parser import HTMLParser
        
        class DocxHTMLParser(HTMLParser):
            def __init__(self, doc):
                super().__init__()
                self.doc = doc
                self.current_para = None
                self.current_run = None
                self.in_code = False
                self.in_pre = False
                
            def handle_starttag(self, tag, attrs):
                if tag == 'h1':
                    self.current_para = self.doc.add_heading('', level=1)
                elif tag == 'h2':
                    self.current_para = self.doc.add_heading('', level=2)
                elif tag == 'h3':
                    self.current_para = self.doc.add_heading('', level=3)
                elif tag == 'h4':
                    self.current_para = self.doc.add_heading('', level=4)
                elif tag == 'p':
                    self.current_para = self.doc.add_paragraph()
                elif tag == 'strong':
                    if self.current_para:
                        self.current_run = self.current_para.add_run()
                        self.current_run.bold = True
                elif tag == 'em':
                    if self.current_para:
                        self.current_run = self.current_para.add_run()
                        self.current_run.italic = True
                elif tag == 'code':
                    self.in_code = True
                elif tag == 'img':
                    # Lấy src từ attrs
                    src = dict(attrs).get('src', '')
                    if src:
                        try:
                            # Download image
                            if src.startswith('/'):
                                # Local file
                                img_path = os.path.join(OUTPUT_FOLDER, src.replace('/outputs/', ''))
                                if os.path.exists(img_path):
                                    self.current_para = self.doc.add_paragraph()
                                    run = self.current_para.add_run()
                                    run.add_picture(img_path, width=Inches(5))
                        except Exception as e:
                            print(f"Error adding image: {e}")
                elif tag == 'ul' or tag == 'ol':
                    self.current_para = self.doc.add_paragraph()
                elif tag == 'li':
                    if self.current_para:
                        self.current_para.style = 'List Bullet' if tag == 'ul' else 'List Number'
                
            def handle_endtag(self, tag):
                if tag in ['h1', 'h2', 'h3', 'h4', 'p']:
                    self.current_para = None
                    self.current_run = None
                elif tag == 'strong' or tag == 'em':
                    self.current_run = None
                elif tag == 'code':
                    self.in_code = False
                    
            def handle_data(self, data):
                if self.in_pre or self.in_code:
                    if self.current_para:
                        run = self.current_para.add_run(data)
                        run.font.name = 'Courier New'
                else:
                    if self.current_run:
                        self.current_run.add_text(data)
                    elif self.current_para:
                        self.current_para.add_run(data)
                    else:
                        self.current_para = self.doc.add_paragraph()
                        self.current_para.add_run(data)
        
        parser = DocxHTMLParser(doc)
        parser.feed(html)
        
        # Save to BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
        
    except ImportError:
        return jsonify({'error': 'python-docx chưa được cài đặt. Chạy: pip install python-docx markdown'}), 500
    except Exception as e:
        return jsonify({'error': f'Lỗi khi tạo DOCX: {str(e)}'}), 500

@app.route('/api/ocr', methods=['POST'])
def ocr():
    """Process OCR request"""
    try:
        if model is None or tokenizer is None:
            return jsonify({
                'success': False,
                'error': 'Model chưa được tải. Vui lòng đợi...'
            }), 503
        
        # Check if file is present
        if 'image' not in request.files:
            return jsonify({
                'success': False,
                'error': 'Không tìm thấy file ảnh'
            }), 400
        
        file = request.files['image']
        if file.filename == '':
            return jsonify({
                'success': False,
                'error': 'Chưa chọn file'
            }), 400
        
        if not allowed_file(file.filename):
            return jsonify({
                'success': False,
                'error': f'Định dạng file không được hỗ trợ. Chỉ chấp nhận: {", ".join(ALLOWED_EXTENSIONS)}'
            }), 400
        
        # Get prompt from form
        prompt_text = request.form.get('prompt', '').strip()
        output_format = request.form.get('format', 'markdown')  # markdown, text, full
        
        # Chọn prompt phù hợp với format - Dùng đúng prompt như DeepSeek-OCR công bố
        if not prompt_text:
            if output_format == 'markdown':
                # Prompt chính xác như DeepSeek-OCR công bố
                prompt_text = '<image>\n<|grounding|>Convert the document to markdown.'
            elif output_format == 'full':
                prompt_text = '<image>\n<|grounding|>OCR this image.'
            else:
                prompt_text = '<image>\nFree OCR.'
        
        # Read image
        image_bytes = file.read()
        image = Image.open(io.BytesIO(image_bytes)).convert('RGB')
        
        # Save uploaded file
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        image.save(filepath)
        
        # Process OCR
        output_path = os.path.join(OUTPUT_FOLDER, f"result_{filename}")
        
        # Giảm kích thước ảnh nếu quá lớn để tránh lỗi CUDA
        actual_image_size = IMAGE_SIZE
        if image.size[0] > 2048 or image.size[1] > 2048:
            # Resize ảnh lớn xuống
            max_dim = max(image.size)
            if max_dim > 2048:
                scale = 2048 / max_dim
                new_size = (int(image.size[0] * scale), int(image.size[1] * scale))
                image = image.resize(new_size, Image.Resampling.LANCZOS)
                image.save(filepath)  # Lưu lại ảnh đã resize
                print(f"⚠️  Ảnh quá lớn, đã resize từ {image.size} xuống {new_size}")
        
        try:
            # Thử với image_size nhỏ hơn nếu gặp lỗi
            result = model.infer(
                tokenizer,
                prompt=prompt_text,
                image_file=filepath,
                output_path=output_path,
                base_size=BASE_SIZE,
                image_size=min(actual_image_size, 640),  # Giới hạn tối đa 640
                crop_mode=CROP_MODE,
                save_results=True,
                test_compress=False  # Tắt test_compress để tránh lỗi
            )
        except RuntimeError as e:
            error_str = str(e)
            if "masked_scatter" in error_str or "CUDA" in error_str:
                # Thử lại với image_size nhỏ hơn
                print(f"⚠️  Lỗi CUDA với image_size={actual_image_size}, thử lại với 512...")
                try:
                    result = model.infer(
                        tokenizer,
                        prompt=prompt_text,
                        image_file=filepath,
                        output_path=output_path,
                        base_size=768,
                        image_size=512,
                        crop_mode=CROP_MODE,
                        save_results=True,
                        test_compress=False
                    )
                except Exception as e2:
                    raise Exception(f"Lỗi khi xử lý OCR (đã thử giảm kích thước): {str(e2)}")
            else:
                raise Exception(f"Lỗi khi xử lý OCR: {error_str}")
        except Exception as e:
            raise Exception(f"Lỗi khi xử lý OCR: {str(e)}")
        
        # Read result - try multiple methods
        result_text = ""
        
        # Method 1: Kiểm tra xem output_path là directory hay file
        possible_files = []
        
        # Nếu output_path là directory, tìm file .txt bên trong
        if os.path.isdir(output_path):
            print(f"📁 Output path là directory: {output_path}")
            # Tìm file .mmd (markdown) trước, sau đó .txt
            for f in os.listdir(output_path):
                if f.endswith('.mmd'):
                    possible_files.append(os.path.join(output_path, f))
            # Nếu không có .mmd, tìm .txt
            if not possible_files:
                for f in os.listdir(output_path):
                    if f.endswith('.txt'):
                        possible_files.append(os.path.join(output_path, f))
            # Nếu vẫn không có, lấy tất cả text files (không phải image)
            if not possible_files:
                for f in os.listdir(output_path):
                    filepath_full = os.path.join(output_path, f)
                    if os.path.isfile(filepath_full) and not f.lower().endswith(('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp')):
                        possible_files.append(filepath_full)
        else:
            # Nếu là file, thử các extension
            possible_files = [
                f"{output_path}.mmd",
                f"{output_path}.txt",
                f"{output_path}",
            ]
        
        # Method 2: Tìm file mới nhất trong OUTPUT_FOLDER có chứa filename
        output_dir = OUTPUT_FOLDER
        if os.path.exists(output_dir):
            all_files = []
            for item in os.listdir(output_dir):
                item_path = os.path.join(output_dir, item)
                # Nếu là directory có chứa filename
                if os.path.isdir(item_path) and filename in item:
                    # Tìm file .mmd hoặc .txt trong directory này
                    for f in os.listdir(item_path):
                        if f.endswith('.mmd') or f.endswith('.txt'):
                            filepath_full = os.path.join(item_path, f)
                            if os.path.isfile(filepath_full):
                                all_files.append((filepath_full, os.path.getmtime(filepath_full)))
                # Nếu là file có chứa filename
                elif os.path.isfile(item_path) and (filename in item or "result_" in item):
                    all_files.append((item_path, os.path.getmtime(item_path)))
            
            if all_files:
                # Sắp xếp theo thời gian, lấy file mới nhất
                all_files.sort(key=lambda x: x[1], reverse=True)
                possible_files.insert(0, all_files[0][0])
        
        # Method 3: Thêm các path khác
        possible_files.extend([
            os.path.join(output_dir, f"result_{filename}.txt"),
            os.path.join(output_dir, f"result_{filename}"),
        ])
        
        # Thử đọc từ các file có thể
        for file_path in possible_files:
            if os.path.exists(file_path) and os.path.isfile(file_path):
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read().strip()
                        if content and len(content) > 10:  # Đảm bảo có nội dung
                            result_text = content
                            print(f"✅ Đã đọc kết quả từ: {file_path}")
                            break
                except Exception as e:
                    print(f"⚠️  Không thể đọc file {file_path}: {e}")
                    continue
        
        # Method 2: Try to get from result object
        if not result_text and result is not None:
            if isinstance(result, dict):
                result_text = result.get('text', result.get('result', result.get('output', str(result))))
            elif isinstance(result, str):
                result_text = result
            elif hasattr(result, 'text'):
                result_text = result.text
            elif hasattr(result, 'output'):
                result_text = result.output
            else:
                result_text = str(result) if result else ""
        
        # Method 4: Nếu vẫn không có, tìm trong OUTPUT_FOLDER file/directory mới nhất
        if not result_text and os.path.exists(OUTPUT_FOLDER):
            try:
                all_items = []
                for item in os.listdir(OUTPUT_FOLDER):
                    item_path = os.path.join(OUTPUT_FOLDER, item)
                    if os.path.isfile(item_path):
                        all_items.append((item_path, os.path.getmtime(item_path), 'file'))
                    elif os.path.isdir(item_path):
                        # Tìm file .txt trong directory
                        for f in os.listdir(item_path):
                            if f.endswith('.mmd') or f.endswith('.txt'):
                                filepath_full = os.path.join(item_path, f)
                                if os.path.isfile(filepath_full):
                                    all_items.append((filepath_full, os.path.getmtime(filepath_full), 'file'))
                                    break
                
                if all_items:
                    # Lấy file mới nhất
                    latest_item = max(all_items, key=lambda x: x[1])
                    latest_path = latest_item[0]
                    with open(latest_path, 'r', encoding='utf-8') as f:
                        content = f.read().strip()
                        if content:
                            result_text = content
                            print(f"✅ Đã đọc kết quả từ file mới nhất: {os.path.basename(latest_path)}")
            except Exception as e:
                print(f"⚠️  Không thể đọc file mới nhất: {e}")
        
        # Clean up và format result text - Giữ nguyên output từ DeepSeek-OCR
        if result_text:
            # Chỉ strip đầu cuối, giữ nguyên mọi thứ bên trong
            result_text = result_text.strip()
            
            # Sửa image paths trong markdown để trỏ đúng output directory
            if output_format == 'markdown' and os.path.isdir(output_path):
                # Tìm tên directory để build path
                output_dir_name = os.path.basename(output_path)
                # Sửa image paths: images/0.jpg -> /outputs/{output_dir_name}/0.jpg
                result_text = re.sub(
                    r'!\[([^\]]*)\]\(images/([^\)]+)\)',
                    lambda m: f'![{m.group(1)}](/outputs/{output_dir_name}/{m.group(2)})',
                    result_text
                )
                # Sửa relative paths khác
                result_text = re.sub(
                    r'!\[([^\]]*)\]\(([^/\)]+\.(jpg|jpeg|png|gif|bmp|webp))\)',
                    lambda m: f'![{m.group(1)}](/outputs/{output_dir_name}/{m.group(2)})',
                    result_text
                )
            
            # Format chỉ khi cần (có tags hoặc format khác markdown)
            # Nếu đã là markdown sạch, giữ nguyên 100%
            formatted_text = format_ocr_result(result_text, output_format)
        else:
            formatted_text = "Không tìm thấy kết quả. Vui lòng kiểm tra logs."
            print(f"⚠️  Không tìm thấy kết quả. Output path: {output_path}")
            print(f"⚠️  Files trong OUTPUT_FOLDER: {os.listdir(OUTPUT_FOLDER) if os.path.exists(OUTPUT_FOLDER) else 'Không tồn tại'}")
        
        return jsonify({
            'success': True,
            'text': formatted_text,
            'raw_text': result_text if result_text else "",
            'filename': filename,
            'format': output_format
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Lỗi xử lý: {str(e)}'
        }), 500

@app.route('/api/ocr-base64', methods=['POST'])
def ocr_base64():
    """Process OCR from base64 image"""
    try:
        if model is None or tokenizer is None:
            return jsonify({
                'success': False,
                'error': 'Model chưa được tải. Vui lòng đợi...'
            }), 503
        
        data = request.get_json()
        if not data or 'image' not in data:
            return jsonify({
                'success': False,
                'error': 'Không tìm thấy dữ liệu ảnh'
            }), 400
        
        # Decode base64 image
        image_data = data['image']
        if image_data.startswith('data:image'):
            image_data = image_data.split(',')[1]
        
        image_bytes = base64.b64decode(image_data)
        image = Image.open(io.BytesIO(image_bytes)).convert('RGB')
        
        # Get prompt
        prompt_text = data.get('prompt', '<image>\nFree OCR.')
        if not prompt_text.strip():
            prompt_text = '<image>\nFree OCR.'
        
        # Save temporary file
        import uuid
        temp_filename = f"temp_{uuid.uuid4().hex}.png"
        temp_filepath = os.path.join(app.config['UPLOAD_FOLDER'], temp_filename)
        image.save(temp_filepath)
        
        # Process OCR
        output_path = os.path.join(OUTPUT_FOLDER, f"result_{temp_filename}")
        
        # Giảm kích thước ảnh nếu quá lớn
        if image.size[0] > 2048 or image.size[1] > 2048:
            max_dim = max(image.size)
            if max_dim > 2048:
                scale = 2048 / max_dim
                new_size = (int(image.size[0] * scale), int(image.size[1] * scale))
                image = image.resize(new_size, Image.Resampling.LANCZOS)
                image.save(temp_filepath)
                print(f"⚠️  Ảnh quá lớn, đã resize từ {image.size} xuống {new_size}")
        
        try:
            result = model.infer(
                tokenizer,
                prompt=prompt_text,
                image_file=temp_filepath,
                output_path=output_path,
                base_size=BASE_SIZE,
                image_size=min(IMAGE_SIZE, 640),  # Giới hạn tối đa 640
                crop_mode=CROP_MODE,
                save_results=True,
                test_compress=False  # Tắt test_compress để tránh lỗi
            )
        except RuntimeError as e:
            error_str = str(e)
            if "masked_scatter" in error_str or "CUDA" in error_str:
                print(f"⚠️  Lỗi CUDA, thử lại với image_size nhỏ hơn...")
                try:
                    result = model.infer(
                        tokenizer,
                        prompt=prompt_text,
                        image_file=temp_filepath,
                        output_path=output_path,
                        base_size=768,
                        image_size=512,
                        crop_mode=CROP_MODE,
                        save_results=True,
                        test_compress=False
                    )
                except Exception as e2:
                    raise Exception(f"Lỗi khi xử lý OCR (đã thử giảm kích thước): {str(e2)}")
            else:
                raise Exception(f"Lỗi khi xử lý OCR: {error_str}")
        except Exception as e:
            raise Exception(f"Lỗi khi xử lý OCR: {str(e)}")
        
        # Read result - try multiple methods (same as ocr function)
        result_text = ""
        
        # Method 1: Kiểm tra xem output_path là directory hay file
        possible_files = []
        
        # Nếu output_path là directory, tìm file .txt bên trong
        if os.path.isdir(output_path):
            print(f"📁 Output path là directory: {output_path}")
            # Tìm file .mmd (markdown) trước, sau đó .txt
            for f in os.listdir(output_path):
                if f.endswith('.mmd'):
                    possible_files.append(os.path.join(output_path, f))
            # Nếu không có .mmd, tìm .txt
            if not possible_files:
                for f in os.listdir(output_path):
                    if f.endswith('.txt'):
                        possible_files.append(os.path.join(output_path, f))
            # Nếu vẫn không có, lấy tất cả text files (không phải image)
            if not possible_files:
                for f in os.listdir(output_path):
                    filepath_full = os.path.join(output_path, f)
                    if os.path.isfile(filepath_full) and not f.lower().endswith(('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp')):
                        possible_files.append(filepath_full)
        else:
            # Nếu là file, thử các extension
            possible_files = [
                f"{output_path}.mmd",
                f"{output_path}.txt",
                f"{output_path}",
            ]
        
        # Method 2: Tìm file mới nhất trong OUTPUT_FOLDER có chứa temp_filename
        if os.path.exists(OUTPUT_FOLDER):
            all_files = []
            for item in os.listdir(OUTPUT_FOLDER):
                item_path = os.path.join(OUTPUT_FOLDER, item)
                # Nếu là directory có chứa temp_filename
                if os.path.isdir(item_path) and temp_filename in item:
                    # Tìm file .mmd hoặc .txt trong directory này
                    for f in os.listdir(item_path):
                        if f.endswith('.mmd') or f.endswith('.txt'):
                            filepath_full = os.path.join(item_path, f)
                            if os.path.isfile(filepath_full):
                                all_files.append((filepath_full, os.path.getmtime(filepath_full)))
                # Nếu là file có chứa temp_filename
                elif os.path.isfile(item_path) and (temp_filename in item or "result_" in item):
                    all_files.append((item_path, os.path.getmtime(item_path)))
            
            if all_files:
                all_files.sort(key=lambda x: x[1], reverse=True)
                possible_files.insert(0, all_files[0][0])
        
        # Method 3: Thêm các path khác
        possible_files.extend([
            os.path.join(OUTPUT_FOLDER, f"result_{temp_filename}.txt"),
            os.path.join(OUTPUT_FOLDER, f"result_{temp_filename}"),
        ])
        
        for file_path in possible_files:
            if os.path.exists(file_path) and os.path.isfile(file_path):
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read().strip()
                        if content and len(content) > 10:
                            result_text = content
                            print(f"✅ Đã đọc kết quả từ: {file_path}")
                            break
                except Exception as e:
                    print(f"⚠️  Không thể đọc file {file_path}: {e}")
                    continue
        
        # Method 2: Try result object
        if not result_text and result is not None:
            if isinstance(result, dict):
                result_text = result.get('text', result.get('result', result.get('output', str(result))))
            elif isinstance(result, str):
                result_text = result
            elif hasattr(result, 'text'):
                result_text = result.text
            elif hasattr(result, 'output'):
                result_text = result.output
            else:
                result_text = str(result) if result else ""
        
        # Method 3: Tìm file mới nhất trong OUTPUT_FOLDER
        if not result_text and os.path.exists(OUTPUT_FOLDER):
            try:
                all_items = []
                for item in os.listdir(OUTPUT_FOLDER):
                    item_path = os.path.join(OUTPUT_FOLDER, item)
                    if os.path.isfile(item_path):
                        all_items.append((item_path, os.path.getmtime(item_path), 'file'))
                    elif os.path.isdir(item_path):
                        # Tìm file .txt trong directory
                        for f in os.listdir(item_path):
                            if f.endswith('.mmd') or f.endswith('.txt'):
                                filepath_full = os.path.join(item_path, f)
                                if os.path.isfile(filepath_full):
                                    all_items.append((filepath_full, os.path.getmtime(filepath_full), 'file'))
                                    break
                
                if all_items:
                    latest_item = max(all_items, key=lambda x: x[1])
                    latest_path = latest_item[0]
                    with open(latest_path, 'r', encoding='utf-8') as f:
                        content = f.read().strip()
                        if content:
                            result_text = content
                            print(f"✅ Đã đọc kết quả từ file mới nhất: {os.path.basename(latest_path)}")
            except Exception as e:
                print(f"⚠️  Không thể đọc file mới nhất: {e}")
        
        if result_text:
            result_text = result_text.strip()
        
        # Clean up temp file
        try:
            os.remove(temp_filepath)
        except:
            pass
        
        return jsonify({
            'success': True,
            'text': result_text if result_text else "Không tìm thấy kết quả. Vui lòng kiểm tra logs."
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Lỗi xử lý: {str(e)}'
        }), 500

if __name__ == '__main__':
    print("=" * 50)
    print("Đang khởi tạo DeepSeek-OCR Web Application...")
    print("=" * 50)
    
    # Initialize model
    if init_model():
        print(f"\nServer đang chạy tại: http://{HOST}:{PORT}")
        print("Nhấn Ctrl+C để dừng server\n")
        app.run(host=HOST, port=PORT, debug=DEBUG)
    else:
        print("Không thể khởi tạo model. Vui lòng kiểm tra lại cấu hình.")

